# app.py

import os
# --- START: ADD THESE TWO LINES ---
from dotenv import load_dotenv
load_dotenv() 
import uuid
# import spacy  # Spacy is still used by ATS, so we keep it
import PyPDF2
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_from_directory, g
from functools import wraps

# --- MODIFICATION 1: SWAP API IMPORTS ---
# import google.generativeai as genai  <- REMOVE THIS
from groq import Groq # <- ADD THIS

from gtts import gTTS
from fpdf import FPDF
import smtplib
import ssl
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import re

import json
import psycopg2
import psycopg2.extras 
import psycopg2.errors
from werkzeug.security import check_password_hash, generate_password_hash

from werkzeug.utils import secure_filename
from ats_calculator import AdvancedATSCalculator
from docx import Document


app = Flask(__name__)
app.secret_key = os.urandom(24)

# --- Configuration Section ---
UPLOAD_FOLDER = 'uploads'
AUDIO_FOLDER = 'static/audio'
REPORTS_FOLDER = 'reports'
RESUME_STORAGE_FOLDER = 'resume_storage'
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://postgres:password@localhost:5432/db_name") 
MAX_QUESTIONS = 10
ATS_UPLOAD_FOLDER = 'ats_uploads'
os.makedirs(ATS_UPLOAD_FOLDER, exist_ok=True)
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
ALLOWED_ITA_EXTENSIONS = {'docx', 'pdf'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(AUDIO_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)
os.makedirs(RESUME_STORAGE_FOLDER, exist_ok=True) 
os.makedirs(ATS_UPLOAD_FOLDER, exist_ok=True)

SENDER_EMAIL = "akshataki7905@gmail.com"
SENDER_APP_PASSWORD = "dwct esbd fucy nybp"
RECIPIENT_EMAIL = "sakshamsaxenamoreyeahs@gmail.com" # This will be the administrator/HR email
# --- END: RE-ADD THIS ENTIRE BLOCK ---


# --- MODIFICATION 2: CONFIGURE THE GROQ CLIENT ---
print("Initializing Groq Client...") 
client = Groq(api_key=os.getenv("GROQ_API_KEY")) 
GROQ_MODEL = "llama3-8b-8192"

# --- MODIFICATION 2: CONFIGURE THE GROQ CLIENT ---
# print("Loaded Google API Key:", os.getenv("GOOGLE_API_KEY"))  <- REMOVE THIS
# genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))            <- REMOVE THIS

# NOTE: The 'spacy' related block is unchanged as it's still needed by the ATS calculator.
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")
# --- END MODIFICATIONS ---


# --- DATABASE HELPER FUNCTIONS (Unchanged) ---
def get_db():
    if 'db' not in g:
        g.db = psycopg2.connect(DATABASE_URL)
    return g.db

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.teardown_appcontext
def close_db(e=None):
    db = g.pop('db', None)
    if db is not None:
        db.close()

# --- ADMIN AUTH DECORATOR (Unchanged) ---
def admin_login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'admin_logged_in' not in session:
            return redirect(url_for('admin_login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

def allowed_ita_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_ITA_EXTENSIONS


# --- MODIFICATION 3: CREATE A CENTRALIZED GROQ API HELPER ---
def call_groq_api(prompt, system_prompt=None):
    """
    A robust helper function to call the Groq API.
    Handles creating the client, managing messages, and returning the response text.
    """
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})

    try:
        chat_completion = client.chat.completions.create(
            messages=messages,
            model=GROQ_MODEL,
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        print(f"Error calling Groq API: {e}")
        # Return an error message or a default value that calling functions can handle
        return f"An error occurred during API call: {e}"
# --- END MODIFICATION ---

# --- MODIFICATION 4: REFACTOR THE ANALYSIS FUNCTION TO USE GROQ ---
# In app.py

# This is the corrected HELPER function.
# It should NOT have an @app.route decorator.
# In app.py

def analyze_transcript_with_groq(transcript, job_description):
    """
    Analyzes an interview transcript using Groq for candidate suitability.
    Dynamically adjusts its prompt and robustly parses the JSON response.
    """
    # Dynamically create the context for the prompt based on job description availability
    if job_description and job_description.strip():
        context_instructions = f"""
        **CONTEXT:**
        - **Job Description:** {job_description}
        - **Interview Transcript:** {transcript}

        **ANALYSIS INSTRUCTIONS:**
        1.  Analyze the candidate's responses in the transcript **strictly against the provided Job Description**.
        """
    else:
        context_instructions = f"""
        **CONTEXT:**
        - **Job Description:** Not Provided.
        - **Interview Transcript:** {transcript}

        **ANALYSIS INSTRUCTIONS:**
        1.  **Infer the Job Role:** Your first and most important task is to carefully analyze the **interviewer's questions** to infer the likely job role, seniority level, and key required skills (e.g., 'Senior Azure Data Engineer', 'Cloud Architect', 'Junior Python Developer').
        2.  **Evaluate Against Inferred Role:** Once you have inferred the role, use it as the benchmark to analyze the candidate's suitability based on their answers.
        """
    
    # Define the full prompt with clear instructions for the AI
    full_prompt = f"""
    You are a world-class HR Analyst and Talent Acquisition Specialist. Your task is to provide a critical, insightful, and fair evaluation of the candidate based on the provided materials.

    {context_instructions}

    **EVALUATION CRITERIA (Provide for all analyses):**
    -   **Overall Score (0-100):** A single integer score reflecting the candidate's overall fit for the role (either provided or inferred).
    -   **Verdict:** A single, concise sentence summarizing your recommendation.
    -   **Green Flags:** 3-5 key strengths or positive signs, citing specific examples from the transcript.
    -   **Red Flags:** 3-5 key weaknesses or areas of concern, citing specific examples.
    -   **Skills Assessment:** A list of demonstrated technical and soft skills.
    -   **Detailed Summary:** A professional paragraph summarizing the candidate's performance, justifying your analysis.

    **CRITICAL:** Your final output MUST BE a single, valid JSON object and nothing else. Do not include markdown formatting like ```json or any surrounding text.

    **JSON OUTPUT STRUCTURE:**
    {{
      "overall_score": <integer>,
      "verdict": "<string>",
      "green_flags": ["<string>", "<string>"],
      "red_flags": ["<string>", "<string>"],
      "technical_skills": ["<string>", "<string>"],
      "soft_skills": ["<string>", "<string>"],
      "detailed_summary": "<string>"
    }}
    """
    
    # Initialize response_text to ensure it's available in the except block for logging
    response_text = ""
    try:
        # Call the centralized Groq API helper function
        response_text = call_groq_api(prompt=full_prompt)
        
        # Check if the helper function itself returned a specific API error message
        if "An error occurred during API call" in response_text:
             raise ValueError(f"Groq API call failed: {response_text}")

        # --- Robust JSON Extraction ---
        # Find the starting '{' and the last '}' to handle cases where the AI is "chatty"
        start_index = response_text.find('{')
        end_index = response_text.rfind('}')

        if start_index != -1 and end_index != -1 and end_index > start_index:
            # Extract the substring that contains only the JSON object
            json_substring = response_text[start_index : end_index + 1]
            
            # Parse the extracted JSON string
            analysis_result = json.loads(json_substring)
            return analysis_result
        else:
            # If a valid JSON object can't be found, raise an error to be caught below
            raise json.JSONDecodeError("Could not find a valid JSON object in the AI's response.", response_text, 0)
    
    except Exception as e:
        # This single except block will now catch all possible errors gracefully:
        # - API call failures
        # - Failure to find a JSON object in the response
        # - Invalid JSON formatting within the object
        print(f"--- ITA ANALYSIS ERROR ---")
        print(f"Error details: {e}")
        print(f"Raw response from AI that caused the error:\n---\n{response_text}\n---")
        
        # Return a consistent error dictionary for the Flask route to handle
        return {"error": "The AI analysis failed. This could be due to a malformed response from the AI or an API connectivity issue. Please check the server logs for more details."}

def cleanup_all_interview_artifacts(session_data):
    """
    Deletes ALL files associated with a completed interview session using robust,
    absolute paths. This includes the temporary AI summary, all audio files,
    the original resume, and the final generated PDF report.
    """
    print("--- Running ROBUST cleanup for ALL interview artifacts... ---")
    
    # We will build a simple list of (folder_path, filename) tuples
    artifacts_to_delete = []

    # 1. The AI-generated summary file in the 'uploads' folder
    if session_data.get('resume_file'):
        artifacts_to_delete.append((UPLOAD_FOLDER, session_data.get('resume_file')))

    # 2. All generated audio files from the 'static/audio' folder
    # Note: AUDIO_FOLDER is 'static/audio', so we don't need to add 'static/' again
    for audio_file in session_data.get('session_audio_files', []):
        artifacts_to_delete.append((AUDIO_FOLDER, audio_file))

    # 3. The candidate's original resume from the 'resume_storage' folder
    if session_data.get('original_resume_filename'):
        artifacts_to_delete.append((RESUME_STORAGE_FOLDER, session_data.get('original_resume_filename')))

    # 4. The final PDF report from the 'reports' folder
    if session_data.get('report_filename'):
        artifacts_to_delete.append((REPORTS_FOLDER, session_data.get('report_filename')))

    # --- Loop through the list and delete each file using its absolute path ---
    for folder, filename in artifacts_to_delete:
        if not filename:
            continue # Skip if filename is empty for some reason

        try:
            # THIS IS THE KEY FIX: Construct a full, absolute path from the app's root.
            # This works reliably regardless of where the script is run from.
            # Example: /path/to/your/project/reports/Summary_....pdf
            full_path = os.path.join(app.root_path, folder, filename)
            
            if os.path.exists(full_path):
                os.remove(full_path)
                print(f"SUCCESS: Cleaned up artifact: {full_path}")
            else:
                print(f"SKIPPING: Artifact not found at: {full_path}")
        except Exception as e:
            print(f"ERROR: Could not delete artifact {folder}/{filename}. Reason: {e}")
            
    print("--- Robust cleanup complete. ---")


# In app.py

# --- NEW, UPDATED FUNCTION ---
def create_summary_pdf(candidate_name, candidate_email, candidate_phone, avg_score, perc_score, final_message, interview_details):
    pdf_filename = f"Summary_{candidate_name.replace(' ', '_')}_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join(REPORTS_FOLDER, pdf_filename)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # --- Report Header ---
    pdf.set_font("Arial", 'B', 20)
    pdf.cell(0, 10, f"Interview Summary for {candidate_name}", 0, 1, 'C')
    pdf.ln(5) # Add a little space
    
    # --- START: ADDED CONTACT INFO BLOCK ---
    pdf.set_font("Arial", '', 11) # Use a smaller, non-bold font
    pdf.cell(0, 6, f"Email: {candidate_email}", 0, 1, 'C')
    pdf.cell(0, 6, f"Phone: {candidate_phone}", 0, 1, 'C')
    pdf.ln(8) # Add more space after contact info
    # --- END: ADDED CONTACT INFO BLOCK ---
    
    # --- Overall Results Section ---
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Overall Results", 0, 1, 'L')
    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 8, f"Average Score: {avg_score:.1f}/10", 0, 1, 'L')
    pdf.cell(0, 8, f"Percentage Score: {perc_score:.1f}%", 0, 1, 'L')
    pdf.multi_cell(0, 8, f"Final Verdict: {final_message}", 0, 'L')
    pdf.ln(5)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
    pdf.ln(5)

    # --- Detailed Breakdown Section (The rest of the function is unchanged) ---
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Detailed Interview Breakdown", 0, 1, 'L')
    pdf.ln(2)
    
    if not interview_details:
        pdf.set_font("Arial", 'I', 12)
        pdf.cell(0, 10, "No questions were answered to provide detailed feedback.", 0, 1, 'L')
    else:
        for i, detail in enumerate(interview_details, 1):
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, f"Question {i} - Score: {detail['score']:.1f}/10", 0, 1, 'L')
            
            pdf.set_font("Arial", 'B', 11)
            clean_question = detail['question'].split(":", 1)[-1].strip() if ":" in detail['question'] else detail['question']
            pdf.multi_cell(0, 6, f'Question Asked: "{clean_question}"', 0, 'L')
            
            pdf.set_font("Arial", '', 11)
            pdf.multi_cell(0, 6, f'Candidate Answer: "{detail["answer"]}"', 0, 'L')
            
            pdf.set_font("Arial", 'I', 11)
            pdf.set_text_color(220, 53, 69) # Set text color to a professional red
            
            feedback_text = detail['feedback'] if detail['feedback'] else "No detailed feedback provided."
            pdf.multi_cell(0, 6, f'Feedback Provided: "{feedback_text}"', 0, 'L')
            
            pdf.set_text_color(0, 0, 0) # Reset text color back to black for the next item
            
            pdf.ln(8) # Add space between questions
            
    try:
        pdf.output(filepath)
        return filepath
    except Exception as e:
        print(f"Error occurred while generating PDF report: {e}")
        return None
    
def send_email_with_pdf(recipient_emails, subject, body, pdf_filepath):
    #... (no changes in this function)
    message = MIMEMultipart()
    message["From"] = SENDER_EMAIL
    message["To"] = ", ".join(recipient_emails)  # Join the list for the 'To' header
    message["Subject"] = subject
    message.attach(MIMEText(body, "plain"))
    
    try:
        with open(pdf_filepath, "rb") as pdf_file:
            attach = MIMEApplication(pdf_file.read(), _subtype="pdf")
            attach.add_header("Content-Disposition", f"attachment; filename={os.path.basename(pdf_filepath)}")
            message.attach(attach)
    except FileNotFoundError:
        print(f"Email Error: PDF file not found at {pdf_filepath}")
        return False
    except Exception as e:
        print(f"Email Error: Could not attach PDF. {e}")
        return False
        
    context = ssl.create_default_context()
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
            server.login(SENDER_EMAIL, SENDER_APP_PASSWORD)
            # The sendmail method accepts a list of recipients directly
            server.sendmail(SENDER_EMAIL, recipient_emails, message.as_string())
            print(f"Email sent successfully to: {', '.join(recipient_emails)}")
            return True
    except Exception as e:
        print(f"Error sending email: {e}")
        return False
# --- END UNCHANGED SECTION ---

# --- MODIFICATION 5: REFACTOR FUNCTIONS TO USE THE GROQ HELPER ---
def extract_info_with_groq(text): #<- Renamed & simplified
    prompt = f"Extract the following information from the text below: Full Name, Education, Previous Experience, Useful Skills. Present the extracted information clearly. Text:\n{text}"
    # Use the centralized helper
    return call_groq_api(prompt=prompt)

def text_to_speech(text, filename):
    # This function is unchanged, it does not use a GenAI model
    try:
        tts = gTTS(text=text, lang='en', slow=False)
        filepath = os.path.join(app.root_path, AUDIO_FOLDER, filename)
        tts.save(filepath)
        return True
    except Exception as e:
        print(f"Error converting text to speech: {e}")
        return False

def generate_dashboard_analytics(interview_details):
    # This function now uses the Groq helper
    interview_transcript = ""
    for detail in interview_details:
        interview_transcript += f"Q: {detail['question']}\nA: {detail['answer']}\n\n"
    SKILLS_TO_EVALUATE = ["Communication", "Technical Knowledge", "Problem-Solving", "Relevant Experience", "Proactivity"]
    skills_prompt = f"""
    Based on the interview transcript below, evaluate the candidate's skills on a scale of 1 to 10 for each category: {', '.join(SKILLS_TO_EVALUATE)}.
    Your output MUST BE a single, valid JSON object. Do not include explanations, notes, or markdown formatting like ```json.
    Example:
    {{
      "Communication": 8, "Technical Knowledge": 6.5, "Problem-Solving": 9, "Relevant Experience": 7, "Proactivity": 8
    }}
    Interview Transcript:
    ---
    {interview_transcript}
    ---
    """
    try:
        # Use the centralized helper
        response_text = call_groq_api(prompt=skills_prompt)
        cleaned_response = response_text.strip().replace("```json", "").replace("```", "").strip()
        skills_data = json.loads(cleaned_response)
        for skill in SKILLS_TO_EVALUATE:
            if skill not in skills_data:
                skills_data[skill] = 0
        return skills_data
    except (json.JSONDecodeError, Exception) as e:
        print(f"Error generating or parsing skills analysis from Groq: {e}")
        return {skill: 0 for skill in SKILLS_TO_EVALUATE}

def conduct_interview_step(role_system_prompt, resume_text, interview_history, candidate_answer, question_num):
    # This function now uses the Groq helper for both steps
    score, feedback = None, None
    if candidate_answer.strip() != "":
        last_question = interview_history[-1] if interview_history else "No previous question found."
        eval_prompt = f"You are an interviewer evaluating a candidate's answer.Make sure the evaluation is done strict in such a way that the answer given must cover the knowledge and understanding behind the question.In case if the cadidate gives entirely irrelevant answer or no answer evaluate the score as 0. Provide a score out of 10 and short constructive feedback. Format the output as: a single decimal number for the score (e.g., 8.5) on the first line, and feedback on subsequent lines.\nInterview Question: {last_question}\nCandidate Answer: {candidate_answer}"
        try:
            # Use helper with a system prompt
            eval_response_text = call_groq_api(prompt=eval_prompt, system_prompt=role_system_prompt)
            lines = eval_response_text.strip().splitlines()
            score = float(lines[0])
            feedback = "\n".join(lines[1:]).strip() if len(lines) > 1 else "Good answer."
        except (ValueError, IndexError, Exception):
            score, feedback = 0, "Could not evaluate the answer."

    next_question_prompt = f"Based on the resume and history, ask the next question.\nResume: {resume_text}\nHistory: {''.join(interview_history)}\nAsk question number {question_num}:"
    try:
        # Use helper with a system prompt again
        next_question = call_groq_api(prompt=next_question_prompt, system_prompt=role_system_prompt).strip()
    except Exception as e:
        print("--- DETAILED GROQ ERROR ---")
        print(e)
        print("-------------------------")
        next_question = "An error occurred with the AI. Please refresh."
        
    return next_question, score, feedback
# --- END MODIFICATION ---

# --- Routes (Main application logic) ---
# Most routes are unchanged, but we update the calls to the refactored functions.

@app.route('/', methods=['GET', 'POST'])
def page1():
    #... No changes here
    if request.method == 'POST':
        # Clean up any leftover files from a previous, unfinished session.
        cleanup_all_interview_artifacts(session.copy())
        session.clear()
        
        candidate_name = request.form.get('candidate_name', '').strip()
        candidate_email = request.form.get('candidate_email', '').strip()
        candidate_phone = request.form.get('candidate_phone', '').strip()

        # --- Server-side validation ---
        if not candidate_name:
            flash("Please enter a candidate name.", "danger")
            return redirect(url_for('page1'))

        # Simple regex for email validation
        if not re.match(r"[^@]+@[^@]+\.[^@]+", candidate_email):
            flash("Please enter a valid email address.", "danger")
            return redirect(url_for('page1'))

        # Simple regex for 10-digit phone number validation
        if not re.match(r"^\d{10}$", candidate_phone):
            flash("Please enter a valid 10-digit phone number.", "danger")
            return redirect(url_for('page1'))
            
        session['candidate_name'] = candidate_name
        session['candidate_email'] = candidate_email
        session['candidate_phone'] = candidate_phone
        
        return redirect(url_for('upload_resume'))
        
    return render_template('page1.html')

@app.route('/upload', methods=['GET', 'POST'])
def upload_resume():
    # --- This route has a small change to call the new Groq function ---
    if 'candidate_name' not in session or 'candidate_email' not in session:
        flash("Please enter your full details first.", "warning")
        return redirect(url_for('page1'))
        
    db_conn = get_db()
    with db_conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute('SELECT id, name FROM roles ORDER BY name')
        roles = cur.fetchall()
    
    if request.method == 'POST':
        role_id = request.form.get('role_id')
        uploaded_file = request.files.get('resume_file')
        job_description = request.form.get('job_description', '').strip()

        if not role_id:
            flash("Please select an interview role.", "warning")
            return render_template('page2.html', roles=roles)

        if not uploaded_file and not job_description:
            flash("Please upload a resume or provide a job description.", "warning")
            return render_template('page2.html', roles=roles)
            
        selected_role_name = next((r['name'] for r in roles if str(r['id']) == str(role_id)), "General")
        file_content_for_ai = ""
        original_resume_filename = None

        if uploaded_file and uploaded_file.filename != '':
            if not allowed_file(uploaded_file.filename):
                flash("Invalid file type. Please upload a PDF, DOCX, or TXT file.", "danger")
                return render_template('page2.html', roles=roles)

            original_fn = secure_filename(uploaded_file.filename)
            unique_fn = f"{uuid.uuid4().hex[:12]}_{original_fn}"
            filepath = os.path.join(RESUME_STORAGE_FOLDER, unique_fn)
            uploaded_file.save(filepath)
            original_resume_filename = unique_fn 

            try:
                if unique_fn.lower().endswith('.pdf'):
                    with open(filepath, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text = page.extract_text()
                            if text: file_content_for_ai += text
                elif unique_fn.lower().endswith('.docx'):
                     doc = Document(filepath)
                     file_content_for_ai = "\n".join([para.text for para in doc.paragraphs])
                else:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        file_content_for_ai = f.read()
            except Exception as e:
                flash(f"Error reading the uploaded file: {e}", "danger")
                return redirect(url_for('upload_resume'))
        else:
            file_content_for_ai = job_description
        
        # --- MODIFICATION: CALL THE GROQ FUNCTION ---
        extracted_text = extract_info_with_groq(file_content_for_ai)
        # --- END MODIFICATION ---
        
        temp_ai_file_id = str(uuid.uuid4())
        save_path = os.path.join(UPLOAD_FOLDER, f"{temp_ai_file_id}.txt")
        with open(save_path, 'w', encoding='utf-8') as f: f.write(extracted_text)
        
        session['resume_file'] = f"{temp_ai_file_id}.txt" 
        session['selected_role_id'] = role_id
        session['selected_role_name'] = selected_role_name
        session['original_resume_filename'] = original_resume_filename 

        session['interview_history'] = []
        session['scores'] = []
        session['feedbacks'] = []
        session['full_questions'] = []
        session['candidate_answers'] = []
        session['question_num'] = 1
        session['interview_finished'] = False
        session['session_audio_files'] = []

        return redirect(url_for('interview'))
        
    return render_template('page2.html', roles=roles)


# The remaining routes (interview, ita_analyzer, etc.) will also call the new Groq functions,
# but the rest of their logic is the same. I've included the fully updated versions below.
@app.route('/behavioral-interview/start', methods=['GET'])
def behavioral_interview_start():
    """
    Sets up a new, purely conversational behavioral interview.
    It does not require a resume or a specific role selection from the user.
    """
    if 'candidate_name' not in session:
        flash("Please provide your details before starting a behavioral interview.", "warning")
        return redirect(url_for('page1'))

    # Clean up any files from a previous session to avoid clutter
    cleanup_all_interview_artifacts(session.copy())

    # Define a list of all keys related to a single interview instance.
    # This prevents clearing the whole session (e.g., admin login).
    INTERVIEW_SESSION_KEYS_TO_CLEAR = [
        'is_behavioral_interview', 'system_prompt_override', 'max_questions', 
        'resume_file', 'selected_role_id', 'selected_role_name', 'original_resume_filename',
        'interview_history', 'scores', 'feedbacks', 'full_questions', 'candidate_answers',
        'question_num', 'interview_finished', 'session_audio_files', 'audio_filename',
        'final_avg_score', 'final_perc_score', 'final_message', 
        'structured_interview_details', 'skills_data', 'line_chart_scores', 'report_filename',
        'ats_resume_text', 'ats_job_description'
    ]
    for key in INTERVIEW_SESSION_KEYS_TO_CLEAR:
        session.pop(key, None)

    # Define the specific instructions for the AI for this interview type
    BEHAVIORAL_SYSTEM_PROMPT = """You are a friendly and professional Human Resources (HR) interviewer conducting a behavioral screening interview. Your goal is to assess the candidate's personality, communication skills, and general professional demeanor.

Your Instructions:
1. **Start the interview** with a classic opening question like 'Tell me about yourself' or 'Can you walk me through your background?'.
2. **Ask exactly 7 questions in total.** After the 7th question, do not ask another.
3. **Build upon the candidate's previous answers** to make the conversation feel natural and engaging.
4. Your questions should be open-ended and behavioral in nature. Focus on topics like teamwork, dealing with conflict, problem-solving under pressure, motivation, and career aspirations.
5. **CRITICAL:** Avoid deep technical questions about specific programming languages, frameworks, or technologies. This is not a technical assessment.
6. Keep your questions concise and clear. Maintain a positive and encouraging tone throughout.

Do not refer to a resume, as you have not been provided with one. This is a purely conversational interview."""
    
    # Initialize the session with settings for a behavioral interview
    session['is_behavioral_interview'] = True
    session['system_prompt_override'] = BEHAVIORAL_SYSTEM_PROMPT
    session['max_questions'] = 7
    session['selected_role_name'] = "Behavioral Interview"
    session['question_num'] = 1
    session['interview_finished'] = False
    session['interview_history'] = []
    session['scores'] = []
    session['feedbacks'] = []
    session['full_questions'] = []
    session['candidate_answers'] = []
    session['session_audio_files'] = []

    session.modified = True
    flash("Starting your behavioral interview. Good luck!", "info")
    return redirect(url_for('interview'))
# --- END: NEW ROUTE ---

# In app.py - This is the fully corrected and flexible interview route

# --- START: MODIFIED INTERVIEW ROUTE ---
@app.route('/interview', methods=['GET', 'POST'])
def interview():
    """
    Handles the main interview flow for both resume-based and behavioral interviews.
    """
    if 'candidate_name' not in session:
        flash("Interview session expired or invalid. Please start again.", "warning")
        return redirect(url_for('page1'))

    # Dynamically set max questions from session, or use the default
    max_interview_questions = session.get('max_questions', MAX_QUESTIONS)

    if session.get('interview_finished') and request.method == 'GET':
        return render_template('intro_interview.html',
            interview_finished=True,
            average_score=session.get('final_avg_score', 0),
            percentage_score=session.get('final_perc_score', 0),
            final_message=session.get('final_message', 'Interview complete.'),
            summary_data=zip(session.get('scores', []), session.get('feedbacks', [])),
            candidate_name=session.get('candidate_name', 'Candidate'),
            report_filename=session.get('report_filename')
        )
    
    # Check for a session-based system prompt (for behavioral interviews).
    role_system_prompt = session.get('system_prompt_override')

    # If no override exists, fetch from the database (for standard interviews).
    if not role_system_prompt:
        if 'selected_role_id' not in session:
            flash("Interview session is missing a role. Please start again.", "warning")
            return redirect(url_for('page1'))
        
        db_conn = get_db()
        with db_conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
            cur.execute('SELECT system_prompt FROM roles WHERE id = %s', (session['selected_role_id'],))
            role = cur.fetchone()
        
        if not role:
            flash("The selected interview role no longer exists. Please start over.", "danger")
            session.clear()
            return redirect(url_for('page1'))
        role_system_prompt = role['system_prompt']

    resume_text_for_ai = ""
    if session.get('resume_file'):
        resume_path = os.path.join(UPLOAD_FOLDER, session['resume_file'])
        if os.path.exists(resume_path):
            with open(resume_path, 'r', encoding='utf-8') as f:
                resume_text_for_ai = f.read()

    if request.method == 'POST':
        is_finished = 'end_interview' in request.form or session.get('question_num', 1) > max_interview_questions
        candidate_answer = request.form.get('answer', '').strip()
        question_num = session.get('question_num', 1)
        interview_history = session.get('interview_history', [])
        
        if candidate_answer:
            _, score, feedback = conduct_interview_step(role_system_prompt, resume_text_for_ai, interview_history, candidate_answer, question_num)
            
            last_question_text = interview_history[-1] if interview_history else "N/A"
            if score is not None:
                session.setdefault('full_questions', []).append(last_question_text)
                session.setdefault('candidate_answers', []).append(candidate_answer)
                session['scores'].append(score)
                session['feedbacks'].append(feedback)
            session['interview_history'].append(f"Answer {question_num-1}: {candidate_answer}")

        if is_finished:
            session['interview_finished'] = True
            
            scores = session.get('scores', [])
            avg_score = sum(scores) / len(scores) if scores else 0
            perc_score = (avg_score / 10) * 100
            final_message = "Interview ended."
            if scores:
                if avg_score > 8: final_message = "Excellent performance! Highly recommended."
                elif avg_score > 6: final_message = "Good performance, a solid candidate."
                elif avg_score > 4: final_message = "Fair performance, with room for improvement."
                else: final_message = "Needs significant improvement."

            structured_interview_details = [{'question': q, 'answer': a, 'score': s, 'feedback': f} for q, a, s, f in zip(session.get('full_questions',[]), session.get('candidate_answers',[]), scores, session.get('feedbacks',[]))]
            skills_data = generate_dashboard_analytics(structured_interview_details)
            
            session['final_avg_score'] = avg_score
            session['final_perc_score'] = perc_score
            session['final_message'] = final_message
            session['structured_interview_details'] = structured_interview_details
            session['skills_data'] = skills_data
            session['line_chart_scores'] = scores

            report_filename = None
            generated_pdf_filepath = create_summary_pdf(
                session['candidate_name'], 
                session.get('candidate_email', 'N/A'),
                session.get('candidate_phone', 'N/A'),
                avg_score, 
                perc_score, 
                final_message, 
                structured_interview_details
            )
            
            if generated_pdf_filepath:
                report_filename = os.path.basename(generated_pdf_filepath)
                session['report_filename'] = report_filename

                all_recipients = [RECIPIENT_EMAIL]
                if session.get('candidate_email'):
                    all_recipients.append(session['candidate_email'])
                
                email_body = f"Please find the interview summary for {session['candidate_name']} attached."
                
                email_sent = send_email_with_pdf(
                    all_recipients, 
                    f"Interview Summary: {session['candidate_name']}", 
                    email_body, 
                    generated_pdf_filepath
                )

                if email_sent:
                    flash(f"Summary PDF for {session['candidate_name']} sent to administrator and candidate.", "success")
                else:
                    flash("Warning: The interview summary PDF could not be sent via email.", "warning")
            else:
                flash("Error: Could not generate the final PDF report.", "danger")

            try:
                db_conn_save = get_db()
                with db_conn_save.cursor() as cur:
                    interview_data_json = json.dumps({ 'details': structured_interview_details, 'skill_analysis': skills_data })
                    cur.execute( """ INSERT INTO interviews (candidate_name, candidate_email, candidate_phone, role_name, final_verdict, average_score, percentage_score, interview_data, resume_filename, report_filename) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s) """, ( session['candidate_name'], session['candidate_email'], session['candidate_phone'], session.get('selected_role_name', 'N/A'), final_message, avg_score, perc_score, interview_data_json, session.get('original_resume_filename'), report_filename ))
                db_conn_save.commit()
                flash("Interview results successfully saved to the database.", "success")
            except Exception as e:
                db_conn_save.rollback()
                print(f"DATABASE ERROR: Failed to save interview results. {e}")
                flash("Error: Could not save the interview results to the database.", "danger")
            
            session_copy_for_cleanup = session.copy()
            # Clear special flags after use
            session.pop('system_prompt_override', None)
            session.pop('is_behavioral_interview', None)
            cleanup_all_interview_artifacts(session_copy_for_cleanup)
            session.modified = True
            
            return render_template('intro_interview.html',
                interview_finished=True, average_score=avg_score, percentage_score=perc_score, final_message=final_message,
                summary_data=zip(scores, session.get('feedbacks', [])), candidate_name=session['candidate_name'], report_filename=report_filename
            )

        next_question, _, _ = conduct_interview_step(role_system_prompt, resume_text_for_ai, session['interview_history'], "", session['question_num'])
        session['interview_history'].append(f"Question {session['question_num']}: {next_question}")
        session['question_num'] += 1
        
        audio_filename = f"q_{session['question_num']-1}_{uuid.uuid4().hex}.mp3"
        if text_to_speech(next_question, audio_filename):
            session['audio_filename'] = audio_filename
            session.setdefault('session_audio_files', []).append(audio_filename)
        current_question = next_question
        
    else: # GET request to start the interview
        session['question_num'] = 1
        first_question, _, _ = conduct_interview_step(role_system_prompt, resume_text_for_ai, [], "", 1)
        session['interview_history'] = [f"Question 1: {first_question}"]
        
        audio_filename = f"q_1_{uuid.uuid4().hex}.mp3"
        if text_to_speech(first_question, audio_filename):
            session['audio_filename'] = audio_filename
            session.setdefault('session_audio_files', []).append(audio_filename)
        current_question = first_question
            
    session.modified = True
    display_question_num = session.get('question_num') if request.method == 'POST' else 1
    
    return render_template('intro_interview.html',
                           question=current_question,
                           last_score=session.get('scores', [])[-1] if session.get('scores') and request.method == 'POST' else None,
                           last_feedback=session.get('feedbacks', [])[-1] if session.get('feedbacks') and request.method == 'POST' else None,
                           question_number=display_question_num,
                           max_questions=max_interview_questions,
                           candidate_name=session.get('candidate_name', 'Candidate'),
                           interview_finished=False,
                           audio_file_url=url_for('static', filename=f'audio/{session.get("audio_filename")}') if session.get('audio_filename') else None)
# --- END: MODIFIED INTERVIEW ROUTE ---

@app.route('/download_report/<path:filename>')
def download_report(filename):
    # No changes here
    if 'report_filename' not in session or session['report_filename'] != filename:
        flash("No report available for download or permission denied.", "danger")
        return redirect(url_for('page1'))
    try:
        return send_from_directory(REPORTS_FOLDER, filename, as_attachment=True)
    except FileNotFoundError:
        flash("Report file not found. It may have been cleaned up.", "warning")
        return redirect(url_for('page1'))

@app.route('/dashboard')
def show_dashboard():
    # No changes here
    if not session.get('interview_finished'):
        flash("You must complete an interview to view the dashboard.", "warning")
        return redirect(url_for('page1'))

    return render_template('dashboard.html',
        candidate_name=session.get('candidate_name', 'N/A'),
        average_score=session.get('final_avg_score', 0),
        percentage_score=session.get('final_perc_score', 0),
        final_verdict=session.get('final_message', 'N/A'),
        skills_data=session.get('skills_data', {}),
        line_chart_scores=session.get('line_chart_scores', []),
        interview_details=session.get('structured_interview_details', [])
    )

# ... The rest of the file (admin routes, ATS checker, etc.) has calls to
# refactored functions, so I'm including the full, corrected code below.

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        db_conn = get_db()
        user = None
        with db_conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
            cur.execute('SELECT * FROM users WHERE username = %s', (username,))
            user = cur.fetchone()
        
        if user and check_password_hash(user['password_hash'], password):
            session['admin_logged_in'] = True
            session['admin_username'] = user['username']
            flash('Login successful!', 'success')
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid username or password', 'danger')
    return render_template('admin_login.html')

@app.route('/admin/dashboard', methods=['GET', 'POST'])
@admin_login_required
def admin_dashboard():
    db_conn = get_db()
    if request.method == 'POST':
        role_name = request.form['name'].strip()
        system_prompt = request.form['system_prompt'].strip()
        if role_name and system_prompt:
            try:
                with db_conn.cursor() as cur:
                    cur.execute('INSERT INTO roles (name, system_prompt) VALUES (%s, %s)', (role_name, system_prompt))
                db_conn.commit()
                flash(f'Role "{role_name}" added successfully.', 'success')
            except psycopg2.errors.UniqueViolation:
                db_conn.rollback()
                flash(f'Error: Role "{role_name}" already exists.', 'danger')
        else:
            flash('Both fields are required.', 'warning')
        return redirect(url_for('admin_dashboard'))
    
    with db_conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute('SELECT * FROM roles ORDER BY name')
        roles = cur.fetchall()
        
    return render_template('admin_dashboard.html', roles=roles)
    
@app.route('/admin/interviews')
@admin_login_required
def admin_interviews():
    db_conn = get_db()
    with db_conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute("""
            SELECT id, candidate_name, candidate_email, candidate_phone, role_name, final_verdict, average_score, 
                   percentage_score, resume_filename, report_filename, created_at
            FROM interviews
            ORDER BY created_at DESC
        """)
        interviews = cur.fetchall()
    return render_template('admin_interviews.html', interviews=interviews)

@app.route('/admin/download_resume/<path:filename>')
@admin_login_required
def admin_download_resume(filename):
    """
    Allows a logged-in admin to download an original resume from the `resume_storage` folder.
    """
    try:
        return send_from_directory(RESUME_STORAGE_FOLDER, filename, as_attachment=True)
    except FileNotFoundError:
        flash(f"Resume file '{filename}' not found on the server.", "danger")
        return redirect(url_for('admin_interviews'))

@app.route('/admin/download_interview_report/<path:filename>')
@admin_login_required
def admin_download_interview_report(filename):
    """
    Allows a logged-in admin to download a PDF report from the persistent 'reports' folder.
    """
    try:
        return send_from_directory(REPORTS_FOLDER, filename, as_attachment=True)
    except FileNotFoundError:
        flash(f"Report file '{filename}' not found on the server.", "danger")
        return redirect(url_for('admin_interviews'))

@app.route('/admin/logout')
@admin_login_required
def admin_logout():
    session.pop('admin_logged_in', None)
    session.pop('admin_username', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('admin_login'))

@app.route('/ats_checker', methods=['GET', 'POST'])
def ats_checker():
    # The AdvancedATSCalculator also uses GenAI, so we need to update its code.
    # The route itself remains the same. The change is inside the calculator's methods.
    # For now, let's assume its `generate_ai_suggestions` is also refactored.
    if request.method == 'POST':
        if 'resume_file' not in request.files:
            flash('No file part in the request.', 'danger')
            return redirect(request.url)
        file = request.files['resume_file']
        job_description = request.form.get('job_description', '').strip()

        if file.filename == '':
            flash('No selected file. Please upload your resume.', 'warning')
            return redirect(request.url)
            
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(ATS_UPLOAD_FOLDER, filename)
            file.save(filepath)
            
            # This class now needs to be initialized with our Groq helper or client
            # The simplest way is to modify its internal call to use our helper.
            calculator = AdvancedATSCalculator() # Make sure to modify AdvancedATSCalculator to use Groq.
            text = calculator.extract_text(filepath, filename)

            if not text:
                os.remove(filepath)
                flash(f"Could not extract text from '{filename}'. The file might be corrupted or empty.", 'danger')
                return redirect(request.url)
            
            results = calculator.calculate_overall_score(text, filename, job_description)
            recommendations = calculator.generate_recommendations(results)
            # This now calls Groq internally within the class method
            ai_suggestions = calculator.generate_ai_suggestions(results, job_description)
            highlighted_resume = calculator.generate_highlighted_resume_html(results)
            
            fig_gauge, fig_radar = calculator.create_visualizations(results)
            
            gauge_html = fig_gauge.to_html(full_html=False, include_plotlyjs='cdn', config={'displayModeBar': False})
            radar_html = fig_radar.to_html(full_html=False, include_plotlyjs=False, config={'displayModeBar': False})

            os.remove(filepath)

            session['ats_resume_text'] = text
            session['ats_job_description'] = job_description
            if not session.get('candidate_name'):
                name_pattern = re.search(r'^([A-Z][a-z]+ [A-Z][a-z]+)', text)
                session['candidate_name'] = name_pattern.group(0) if name_pattern else "Candidate"
            
            session.modified = True

            return render_template(
                'ats_checker.html',
                results=results,
                recommendations=recommendations,
                ai_suggestions=ai_suggestions,
                highlighted_resume=highlighted_resume,
                gauge_chart=gauge_html,
                radar_chart=radar_html
            )
        else:
            flash('Invalid file type. Please upload a PDF, DOCX, or TXT file.', 'danger')
            return redirect(request.url)

    return render_template('ats_checker.html', results=None)

@app.route('/interview/start_from_ats', methods=['GET'])
def start_interview_from_ats():
    # Calls extract_info_with_groq now
    if 'ats_resume_text' not in session:
        flash("Please analyze your resume first before starting an interview.", "warning")
        return redirect(url_for('ats_checker'))

    cleanup_all_interview_artifacts(session.copy())
    
    resume_text = session['ats_resume_text']
    job_description = session['ats_job_description']

    db_conn = get_db()
    role_id, role_name = None, None
    with db_conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute("SELECT id, name FROM roles WHERE name ILIKE '%Software%' ORDER BY id LIMIT 1")
        role = cur.fetchone()
        if not role:
            cur.execute("SELECT id, name FROM roles ORDER BY id LIMIT 1")
            role = cur.fetchone()
        if role:
            role_id, role_name = role['id'], role['name']

    if not role_id:
        flash("No suitable interview roles are configured. Please ask an admin to add one.", "danger")
        return redirect(url_for('page1'))

    combined_info = f"--- RESUME TEXT ---\n{resume_text}\n\n--- JOB DESCRIPTION ---\n{job_description}"
    
    original_resume_filename = f"ats_generated_{uuid.uuid4().hex[:12]}.txt"
    filepath = os.path.join(RESUME_STORAGE_FOLDER, original_resume_filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(combined_info)
    
    # MODIFIED CALL
    extracted_text = extract_info_with_groq(combined_info)
    
    temp_ai_file_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_FOLDER, f"{temp_ai_file_id}.txt")
    with open(save_path, 'w', encoding='utf-8') as f:
        f.write(extracted_text)

    session['resume_file'] = f"{temp_ai_file_id}.txt"
    session['selected_role_id'] = role_id
    session['selected_role_name'] = role_name
    session['original_resume_filename'] = original_resume_filename
    
    session['interview_history'] = []
    session['scores'] = []
    session['feedbacks'] = []
    session['full_questions'] = []
    session['candidate_answers'] = []
    session['question_num'] = 1
    session['interview_finished'] = False
    session['session_audio_files'] = []

    session.pop('ats_resume_text', None)
    session.pop('ats_job_description', None)

    flash("Great! Let's start the interview for this role.", "success")
    return redirect(url_for('interview'))


@app.route('/ita_analyzer', methods=['GET', 'POST'])
def ita_analyzer():
    if request.method == 'POST':
        if 'transcript_file' not in request.files:
            flash('No transcript file part in the request.', 'danger')
            return redirect(request.url)
            
        file = request.files['transcript_file']
        job_description = request.form.get('job_description', '').strip()

        if file.filename == '':
            flash('No selected file. Please upload a transcript.', 'warning')
            return redirect(request.url)
        
        if file and allowed_ita_file(file.filename):
            try:
                transcript_text = ""
                if file.filename.lower().endswith('.docx'):
                    doc = Document(file)
                    transcript_text = "\n".join([para.text for para in doc.paragraphs])
                elif file.filename.lower().endswith('.pdf'):
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        transcript_text += page.extract_text() or ""
                
                if not transcript_text.strip():
                    flash('The uploaded file appears to be empty or could not be read.', 'danger')
                    return redirect(request.url)
                
                # Correctly calling the helper function with required arguments
                results = analyze_transcript_with_groq(transcript_text, job_description)

                # Correctly checking for an error dictionary from the helper
                if "error" in results:
                    flash(results["error"], 'danger')
                    return render_template('ita_analyzer.html', results=None)

                # If no error, render the results
                return render_template('ita_analyzer.html', results=results)

            except Exception as e:
                # Catch any other unexpected errors during file processing
                flash(f'An error occurred while processing the file: {e}', 'danger')
                return redirect(request.url)
        else:
            flash('Invalid file type. Please upload a DOCX or PDF file.', 'danger')
            return redirect(request.url)
            
    # For a GET request, just show the upload form
    return render_template('ita_analyzer.html', results=None)

if __name__ == '__main__':
    app.run(debug=True)
